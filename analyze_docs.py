#!/usr/bin/env python3
"""
Document Analysis System
Ingests, digests, and suggests paths based on Word documents in the repository.
"""

import re
import sys
from pathlib import Path
from typing import List, Dict, Any
from docx import Document


class DocumentAnalyzer:
    """Handles document ingestion, digestion, and path suggestion."""
    
    def __init__(self, base_dir: str = "."):
        self.base_dir = Path(base_dir)
        if not self.base_dir.exists() or not self.base_dir.is_dir():
            print(f"Error: The base directory '{self.base_dir}' does not exist or is not a directory.")
            sys.exit(1)
        self.documents: Dict[str, Any] = {}
        
    def ingest_docs(self) -> List[str]:
        """
        Ingest all .docx files from the repository.
        Returns a list of ingested document names.
        """
        print("=" * 80)
        print("INGESTING DOCUMENTS")
        print("=" * 80)
        
        docx_files = [
            f for f in self.base_dir.glob("**/*.docx")
            if not any(part.startswith('.') or part.startswith('~$') for part in f.parts)
        ]
        docx_files = list(self.base_dir.glob("*.docx"))
        ingested = []
        
        for docx_file in docx_files:
            try:
                doc = Document(docx_file)
                self.documents[docx_file.name] = {
                    'path': docx_file,
                    'document': doc,
                    'paragraphs': [p.text for p in doc.paragraphs if p.text.strip()],
                    'num_paragraphs': len([p for p in doc.paragraphs if p.text.strip()])
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                self.documents[docx_file.name] = {
                    'path': docx_file,
                    'paragraphs': paragraphs,
                    'num_paragraphs': len(paragraphs)
                }
                ingested.append(docx_file.name)
                print(f"✓ Ingested: {docx_file.name}")
            except Exception as e:
                print(f"✗ Failed to ingest {docx_file.name}: {e}")
                print(f"✗ Failed to ingest {docx_file.name}: {type(e).__name__}: {e}")
        
        print(f"\nTotal documents ingested: {len(ingested)}")
        return ingested
    
    def digest_docs(self) -> Dict[str, Any]:
        """
        Digest the ingested documents by analyzing their content.
        Returns a summary of the analysis.
        
        Returns:
            Dict[str, Any]: A summary of the analysis with the following structure:
                {
                    'total_documents': int,  # Total number of ingested documents
                    'documents': {
                        <doc_name>: {
                            'num_paragraphs': int,      # Number of paragraphs in the document
                            'word_count': int,          # Total word count in the document
                            'key_topics': List[str],    # List of key topics extracted from the document
                            'summary': str              # Brief summary from the first few paragraphs
                        },
                        ...
                    }
                }
        """
        print("\n" + "=" * 80)
        print("DIGESTING DOCUMENTS")
        print("=" * 80)
        
        analysis = {
            'total_documents': len(self.documents),
            'documents': {}
        }
        
        for doc_name, doc_data in self.documents.items():
            paragraphs = doc_data['paragraphs']
            text = ' '.join(paragraphs)
            
            # Extract key information
            doc_analysis = {
                'num_paragraphs': doc_data['num_paragraphs'],
                'word_count': len(re.findall(r'\b\w+\b', text)),
                'key_topics': self._extract_key_topics(text),
                'word_count': sum(len(p.split()) for p in paragraphs),
                'key_topics': self._extract_key_topics(paragraphs),
                'summary': self._create_summary(paragraphs[:3])  # First 3 paragraphs
            }
            
            analysis['documents'][doc_name] = doc_analysis
            
            print(f"\n📄 {doc_name}")
            print(f"   Paragraphs: {doc_analysis['num_paragraphs']}")
            print(f"   Word Count: {doc_analysis['word_count']}")
            print(f"   Key Topics: {', '.join(doc_analysis['key_topics'][:5])}")
        
        return analysis
    
    def _extract_key_topics(self, text: str) -> List[str]:
        """Extract key topics from text using keyword frequency."""
        # Convert to lowercase and extract words using regex
        words = re.findall(r'\b\w+\b', text.lower())
        
        # Filter common words and short words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                      'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
                      'this', 'that', 'these', 'those', 'it', 'its', 'will', 'can', 'may',
                      'would', 'could', 'should', 'has', 'have', 'had', 'do', 'does', 'did'}
        
        # Count word frequencies
        word_freq = {}
        for word in words:
            if len(word) > 3 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1
            if doc_analysis['summary']:
                print(f"   Summary: {doc_analysis['summary']}")
        
        return analysis
    
    def _extract_key_topics(self, paragraphs: List[str]) -> List[str]:
        """
        Extract key topics from text using simple frequency analysis.

        Algorithm:
        - Converts text to lowercase and splits into words.
        - Removes punctuation from each word.
        - Applies basic stemming (removes -ing, -ed, -s suffixes).
        - Filters out stop words and words with 3 or fewer characters.
        - Counts the frequency of remaining words.
        - Returns the top 10 most frequent words as key topics.

        Limitations:
        - Only considers word frequency; does not use semantic analysis.
        - May miss multi-word topics or context-specific terms.
        - Basic stemming may incorrectly modify some words.
        - The stop word list is static and may not cover all common words.
        
        Args:
            paragraphs (List[str]): List of paragraph texts to analyze.
            
        Returns:
            List[str]: Top 10 most frequent words as key topics.
        """
        # Expanded stop words list for better filtering
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
            'this', 'that', 'these', 'those', 'it', 'its', 'will', 'can', 'may',
            'would', 'could', 'should', 'has', 'have', 'had', 'do', 'does', 'did',
            'they', 'their', 'them', 'we', 'you', 'your', 'our', 'who', 'which',
            'what', 'when', 'where', 'how', 'all', 'each', 'some', 'more', 'most',
            'other', 'into', 'through', 'during', 'before', 'after', 'above',
            'below', 'between', 'under', 'again', 'further', 'then', 'once',
            'here', 'there', 'than', 'such', 'only', 'very', 'just', 'also',
            'being', 'both', 'about', 'over', 'any', 'same', 'own', 'while'
        }
        
        # Count word frequencies across all paragraphs
        word_freq = {}
        for paragraph in paragraphs:
            words = paragraph.lower().split()
            for word in words:
                # Remove punctuation
                clean_word = ''.join(c for c in word if c.isalnum())
                # Apply basic stemming (remove common suffixes)
                if clean_word.endswith('ing'):
                    clean_word = clean_word[:-3]
                elif clean_word.endswith('ed'):
                    clean_word = clean_word[:-2]
                elif clean_word.endswith('s') and len(clean_word) > 4:
                    clean_word = clean_word[:-1]
                
                if len(clean_word) > 3 and clean_word not in stop_words:
                    word_freq[clean_word] = word_freq.get(clean_word, 0) + 1
        
        # Sort by frequency and return top words
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:10]]
    
    def _create_summary(self, paragraphs: List[str]) -> str:
        """Create a brief summary from the first few paragraphs."""
        summary = ' '.join(paragraphs)
        if len(summary) > 300:
            summary = summary[:297] + "..."
            truncated = summary[:297]
            # Truncate at the last complete word before the cutoff
            if ' ' in truncated:
                truncated = truncated.rsplit(' ', 1)[0]
            summary = truncated + "..."
        return summary
    
    def suggest_path(self, analysis: Dict[str, Any]) -> None:
        """
        Suggest a development path based on the document analysis.
        
        Args:
            analysis (Dict[str, Any]): Dictionary containing document analysis results.
                The dictionary must include a 'documents' key mapping to per-document statistics,
                where each value is a dict with keys such as 'key_topics', 'summary', etc.
        """
        print("\n" + "=" * 80)
        print("SUGGESTED PATH")
        print("=" * 80)
        
        print("\nBased on the analysis of the ingested documents, here are the recommendations:\n")
        
        # Categorize documents by content
        fitness_docs = []
        business_docs = []
        gamification_docs = []
        
        for doc_name, doc_data in analysis['documents'].items():
            topics = [t.lower() for t in doc_data['key_topics']]
            
            if any(word in topics for word in ['fitness', 'training', 'workout', 'exercise', 'gym']):
                fitness_docs.append(doc_name)
            if any(word in topics for word in ['business', 'enterprise', 'coaching', 'model']):
                business_docs.append(doc_name)
            if any(word in topics for word in ['gamified', 'game', 'battle', 'quest']):
                gamification_docs.append(doc_name)
        
        # Show document categorization first
        print("=" * 80)
        print("DOCUMENT CATEGORIZATION")
        print("=" * 80)
        
        if fitness_docs:
            print(f"\n💪 Fitness-focused documents ({len(fitness_docs)}):")
            for doc in fitness_docs:
                print(f"   - {doc}")
        
        if business_docs:
            print(f"\n💼 Business-focused documents ({len(business_docs)}):")
            for doc in business_docs:
                print(f"   - {doc}")
        
        if gamification_docs:
            print(f"\n🎮 Gamification-focused documents ({len(gamification_docs)}):")
            for doc in gamification_docs:
                print(f"   - {doc}")
        
        # Generate dynamic development path based on analysis
        print("\n" + "=" * 80)
        print("📋 RECOMMENDED DEVELOPMENT PATH (Generated from Analysis)")
        print("=" * 80)
        print()
        
        phase = 1
        
        # Always start with foundation if there are any documents
        if fitness_docs or business_docs or gamification_docs:
            print(f"{phase}. BUILD THE FOUNDATION")
            if fitness_docs:
                print("   └─ Create core data models for fitness coaching")
                print("   └─ Define user profiles and progress tracking structures")
            if business_docs:
                print("   └─ Establish client management and business workflows")
            if gamification_docs:
                print("   └─ Set up gamification mechanics (points, levels, achievements)")
            phase += 1
            print()
        
        # Add interface development if we have fitness or business content
        if fitness_docs or business_docs:
            print(f"{phase}. DEVELOP THE INTERFACE")
            if fitness_docs:
                print("   └─ Design user dashboard for fitness tracking")
                print("   └─ Create workout logging interface")
                print("   └─ Build progress visualization tools")
            if business_docs:
                print("   └─ Develop coach management dashboard")
                print("   └─ Create client onboarding interface")
            phase += 1
            print()
        
        # Add coaching features if we have fitness content
        if fitness_docs:
            print(f"{phase}. IMPLEMENT COACHING FEATURES")
            print("   └─ Personalized workout recommendations")
            print("   └─ Goal setting and tracking")
            print("   └─ Motivational messaging system")
            print("   └─ Progress assessment and feedback")
            phase += 1
            print()
        
        # Add gamification layer if we detected gamification focus
        if gamification_docs:
            print(f"{phase}. ADD GAMIFICATION LAYER")
            print("   └─ Achievement system based on milestones")
            print("   └─ Challenge modes and battle scenarios")
            print("   └─ Leaderboards and social features")
            print("   └─ Reward mechanisms and badges")
            phase += 1
            print()
        else:
            # Suggest gamification as optional if not strongly detected
            if fitness_docs or business_docs:
                print(f"{phase}. CONSIDER GAMIFICATION (Optional)")
                print("   └─ NOTE: No strong gamification signals detected in documents")
                print("   └─ Consider if gamification aligns with business goals")
                print("   └─ Could add achievement system for user engagement")
                phase += 1
                print()
        
        # Add business integration if we have business content
        if business_docs:
            print(f"{phase}. BUSINESS INTEGRATION")
            print("   └─ Payment and subscription management")
            print("   └─ Coach-client communication tools")
            print("   └─ Analytics and reporting dashboard")
            print("   └─ Marketing and customer acquisition features")
            phase += 1
            print()
        
        print("=" * 80)
        print("NEXT STEPS")
        print("=" * 80)
        
        # Generate dynamic next steps based on what was found
        next_steps = []
        
        if analysis['total_documents'] > 0:
            next_steps.append("Review the ingested document content in detail")
            
        if fitness_docs:
            next_steps.append("Define fitness coaching data models and workout structures")
            
        if gamification_docs:
            next_steps.append("Design gamification mechanics and reward systems")
            
        if business_docs:
            next_steps.append("Outline business requirements and revenue models")
            
        next_steps.extend([
            "Create technical specifications for each component",
            "Set up project structure (frontend, backend, database)",
            "Begin iterative development with MVP features first"
        ])
        
        for i, step in enumerate(next_steps, 1):
            print(f"\n{i}. {step}")
            # Check for fitness-related terms (considering stemming)
            if any(word in topics for word in ['fitnes', 'fitness', 'train', 'training', 'workout', 'exercise', 'gym']):
                fitness_docs.append(doc_name)
            # Check for business-related terms (considering stemming)
            if any(word in topics for word in ['busines', 'business', 'enterprise', 'coach', 'coaching', 'model']):
                business_docs.append(doc_name)
            # Check for gamification-related terms (considering stemming)
            if any(word in topics for word in ['gamifi', 'gamified', 'game', 'battl', 'battle', 'quest']):
                gamification_docs.append(doc_name)
        
        print("📋 RECOMMENDED DEVELOPMENT PATH:\n")
        
        # Build dynamic path based on detected themes
        if fitness_docs or business_docs or gamification_docs:
            print("1. BUILD THE FOUNDATION")
            print("   └─ Create core data models for fitness coaching")
            print("   └─ Define user profiles and progress tracking structures")
            if gamification_docs:
                print("   └─ Establish gamification mechanics (points, levels, achievements)")
            
            if fitness_docs:
                print("\n2. DEVELOP THE INTERFACE")
                print("   └─ Design user dashboard with gamified elements")
                print("   └─ Create workout logging interface")
                print("   └─ Build progress visualization tools")
                
                print("\n3. IMPLEMENT COACHING FEATURES")
                print("   └─ Personalized workout recommendations")
                print("   └─ Goal setting and tracking")
                print("   └─ Motivational messaging system")
            
            if gamification_docs:
                print("\n4. ADD GAMIFICATION LAYER")
                print("   └─ Achievement system based on milestones")
                print("   └─ Challenge modes and battle scenarios")
                print("   └─ Leaderboards and social features")
            
            if business_docs:
                print("\n5. BUSINESS INTEGRATION")
                print("   └─ Payment and subscription management")
                print("   └─ Coach-client communication tools")
                print("   └─ Analytics and reporting dashboard")
        else:
            # Default roadmap if no specific themes detected
            print("1. BUILD THE FOUNDATION")
            print("   └─ Create core data models for fitness coaching")
            print("   └─ Define user profiles and progress tracking structures")
            print("   └─ Establish gamification mechanics (points, levels, achievements)")
            
            print("\n2. DEVELOP THE INTERFACE")
            print("   └─ Design user dashboard with gamified elements")
            print("   └─ Create workout logging interface")
            print("   └─ Build progress visualization tools")
            
            print("\n3. IMPLEMENT COACHING FEATURES")
            print("   └─ Personalized workout recommendations")
            print("   └─ Goal setting and tracking")
            print("   └─ Motivational messaging system")
            
            print("\n4. ADD GAMIFICATION LAYER")
            print("   └─ Achievement system based on milestones")
            print("   └─ Challenge modes and battle scenarios")
            print("   └─ Leaderboards and social features")
            
            print("\n5. BUSINESS INTEGRATION")
            print("   └─ Payment and subscription management")
            print("   └─ Coach-client communication tools")
            print("   └─ Analytics and reporting dashboard")
        
        print("\n" + "=" * 80)
        print("DOCUMENT CATEGORIZATION")
        print("=" * 80)
        
        if fitness_docs:
            print(f"\n💪 Fitness-focused documents ({len(fitness_docs)}):")
            for doc in fitness_docs:
                print(f"   - {doc}")
        
        if business_docs:
            print(f"\n💼 Business-focused documents ({len(business_docs)}):")
            for doc in business_docs:
                print(f"   - {doc}")
        
        if gamification_docs:
            print(f"\n🎮 Gamification-focused documents ({len(gamification_docs)}):")
            for doc in gamification_docs:
                print(f"   - {doc}")
        
        print("\n" + "=" * 80)
        print("NEXT STEPS")
        print("=" * 80)
        print("\n1. Review the ingested document content in detail")
        print("2. Prioritize features based on business requirements")
        print("3. Create technical specifications for each component")
        print("4. Set up project structure (frontend, backend, database)")
        print("5. Begin iterative development with MVP features first")
        
        print("\n✓ Analysis complete!")
    
    def run(self) -> None:
        """Execute the full pipeline: ingest, digest, suggest."""
        ingested = self.ingest_docs()
        
        if not ingested:
            print("\n⚠ No documents found to analyze!")
            return
        
        analysis = self.digest_docs()
        self.suggest_path(analysis)


def main():
    """Main entry point for the document analyzer."""
    # Get the script directory or use current directory
    base_dir = Path(__file__).parent if '__file__' in globals() else Path.cwd()
    
    analyzer = DocumentAnalyzer(base_dir=base_dir)
    analyzer.run()


if __name__ == "__main__":
    main()
